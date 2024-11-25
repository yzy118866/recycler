import logging
import os
import re

from docx.api import Document
from rest_framework.exceptions import APIException

from recycler.models import FKKO, CompanyContract, Company

log = logging.getLogger("parse_fkko")


def extract_numbers(text: str) -> str:
    return re.sub(r"[^\d]*", "", text)


def extract_security_class(text: str) -> str:
    """Extract the security class (roman numerals from I to V)"""
    roman_numerals = {
        "I": 1, "II": 2, "III": 3, "IV": 4, "V": 5
    }
    text = text.strip().upper()
    if text in roman_numerals:
        return text
    return None

def extract_contract_number(document):
    for para in document.paragraphs:
        match = re.search(r"ДОГОВОР\s*№?\s*([A-Za-z0-9/]+)", para.text)
        if match:
            contract_num = match.group(1).strip()
            formatted_contract_num = f"№ {contract_num}"
            log.info(f"Найден номер договора: {formatted_contract_num}")
            return formatted_contract_num
    return None

def parse_fkko(contract: CompanyContract):
    """Parse FKKO list from docx file tables"""
    
    file_path = contract.file.path
    file_ext = os.path.splitext(file_path)[1]

    if file_ext not in [".docx"]:
        raise APIException(f"Неверное расширение файла ({file_ext}). Парсинг ФККО работает только с DOCX.", code=422)

    document = Document(file_path)

    fkko_list: list[FKKO] = []
    fkko_security_classes = {}

    contract_num = extract_contract_number(document)
    if contract_num:
        log.info(f"Номер договора: {contract_num}")

        company = contract.company
        company.contract_num = contract_num
        company.save()

        updated_company = Company.objects.get(pk=company.pk)
        log.info(f"Сохраненное значение номера договора: {updated_company.contract_num}")

    for table in document.tables:
        if len(table.rows) > 0 and len(table.rows[0].cells) == 5:
            for row in table.rows:
                cell_texts = [cell.text.strip() for cell in row.cells]
                if len(cell_texts) == 5:
                    code = extract_numbers(cell_texts[2])
                    security_class = extract_security_class(cell_texts[3])
                    if code and security_class:
                        fkko_security_classes[code] = security_class
                        log.info(f"Класс опасности для кода {code}: {security_class}")

    for table in document.tables:
        for i, row in enumerate(table.rows):
            cell_texts = list(cell.text.strip() for cell in row.cells)
            
            if len(cell_texts) == 7:
                try:
                    name = cell_texts[1]
                    code = extract_numbers(cell_texts[2])
                    price_text = cell_texts[5]

                    price_text_cleaned = (
                        price_text.replace("\u00A0", "")
                        .replace(" ", "")
                        .replace(",", ".")
                        .strip()
                    )

                    price = int(float(price_text_cleaned))

                    if not code or not price or not name:
                        log.warning(f"Пропущена некорректная строка: {cell_texts}")
                        continue

                    security_class = fkko_security_classes.get(code)
                    if not security_class:
                        log.warning(f"Не найден класс опасности для кода {code}. Пропускаем строку.")
                        continue

                    fkko_list.append(FKKO(
                        contract=contract,
                        company=contract.company,
                        code=code,
                        price=price,
                        name=name,
                        security_class=security_class
                    ))

                except Exception as e:
                    log.warning(f"Ошибка обработки строки: {cell_texts}, ошибка: {e}")
                    continue

    res = FKKO.objects.bulk_create(
        fkko_list,
        update_conflicts=True,
        unique_fields=["code", "company"],
        update_fields=["name", "price", "security_class"],
    )
    updated_count = len(res)
    log.info(f"Обновлено записей ФККО: {updated_count}")
    return updated_count
